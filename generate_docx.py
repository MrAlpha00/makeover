import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def add_page_border(section):
    sectPr = section._sectPr
    pgBorders = OxmlElement('w:pgBorders')
    pgBorders.set(qn('w:offsetFrom'), 'page')
    for border_side in ['top', 'bottom', 'left', 'right']:
        border = OxmlElement(f'w:{border_side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '6')
        border.set(qn('w:space'), '24')
        border.set(qn('w:color'), 'auto')
        pgBorders.append(border)
    sectPr.append(pgBorders)

def add_toc(doc):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._r.append(instrText)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)
    doc.add_paragraph()

def set_font_styles(doc):
    normal_style = doc.styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Times New Roman'
    normal_font.size = Pt(12)
    normal_style.paragraph_format.space_after = Pt(6)

    heading1_style = doc.styles['Heading 1']
    h1_font = heading1_style.font
    h1_font.name = 'Times New Roman'
    h1_font.size = Pt(16)
    h1_font.bold = True
    heading1_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading1_style.paragraph_format.space_before = Pt(24)
    heading1_style.paragraph_format.space_after = Pt(12)

    heading2_style = doc.styles['Heading 2']
    h2_font = heading2_style.font
    h2_font.name = 'Times New Roman'
    h2_font.size = Pt(14)
    h2_font.bold = True
    heading2_style.paragraph_format.space_before = Pt(12)
    heading2_style.paragraph_format.space_after = Pt(6)

    heading3_style = doc.styles['Heading 3']
    h3_font = heading3_style.font
    h3_font.name = 'Times New Roman'
    h3_font.size = Pt(12)
    h3_font.bold = True
    h3_font.italic = True
    heading3_style.paragraph_format.space_before = Pt(6)
    heading3_style.paragraph_format.space_after = Pt(4)

def add_header_footer(section, doc):
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = "PartyHubs Final Year Project Report"
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_font = header_para.style.font
    header_font.name = 'Times New Roman'
    header_font.size = Pt(10)

    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    run._r.append(instrText)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)
    footer_font = footer_para.style.font
    footer_font.name = 'Times New Roman'
    footer_font.size = Pt(10)

def parse_markdown(doc, md_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.read().split('\n')

    in_code_block = False
    code_lines = []
    in_table = False
    table_rows = []
    prev_line_empty = False

    for line in lines:
        stripped = line.strip()

        if stripped.startswith('!['):
            doc.add_paragraph(f'[Image Placeholder: {stripped}]', style='Normal')
            continue

        if stripped.startswith('# '):
            text = stripped[2:].strip()
            if 'Chapter ' in text and 'Chapter 1' not in text:
                doc.add_page_break()
            if text == 'Table of Contents':
                doc.add_heading(text, level=1)
                add_toc(doc)
            else:
                doc.add_heading(text, level=1)
            continue

        if stripped.startswith('## '):
            text = stripped[3:].strip()
            doc.add_heading(text, level=2)
            continue

        if stripped.startswith('### '):
            text = stripped[4:].strip()
            doc.add_heading(text, level=3)
            continue

        if stripped.startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lines = []
            else:
                in_code_block = False
                code_text = '\n'.join(code_lines)
                para = doc.add_paragraph(code_text)
                para.style.font.name = 'Courier New'
                para.style.font.size = Pt(10)
                para.paragraph_format.left_indent = Cm(1)
                doc.add_paragraph()
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        if stripped.startswith('|'):
            if not in_table:
                in_table = True
                table_rows = []
            cells = [c.strip() for c in stripped.split('|') if c.strip()]
            if cells and all(re.match(r'^-+$', c) for c in cells):
                continue
            table_rows.append(cells)
            continue
        else:
            if in_table:
                in_table = False
                if table_rows:
                    cols = len(table_rows[0])
                    table = doc.add_table(rows=len(table_rows), cols=cols)
                    for i, row in enumerate(table_rows):
                        for j in range(min(len(row), cols)):
                            table.rows[i].cells[j].text = row[j]
                    table.style = 'Light Grid Accent 1'
                    doc.add_paragraph()
                table_rows = []

        if stripped.startswith('- '):
            doc.add_paragraph(stripped[2:].strip(), style='List Bullet')
            continue

        if not stripped:
            if not prev_line_empty:
                doc.add_paragraph()
            prev_line_empty = True
            continue
        else:
            prev_line_empty = False
            doc.add_paragraph(stripped, style='Normal')

def main():
    md_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'report.md')
    docx_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'report.docx')

    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found")
        return

    doc = Document()

    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    add_page_border(section)
    set_font_styles(doc)
    add_header_footer(section, doc)
    parse_markdown(doc, md_path)

    doc.save(docx_path)
    print(f"Generated {docx_path}")
    return docx_path

if __name__ == '__main__':
    main()
