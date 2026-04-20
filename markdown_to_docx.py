import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

input_file = "party-hub-content.md"
output_file = "party-hub-content.docx"

def convert_markdown_to_docx(md_file, docx_file):
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()

    lines = content.split('\n')
    
    for line in lines:
        line = line.rstrip()
        
        if line.startswith('# '):
            heading = doc.add_heading(line[2:], 0)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in heading.runs:
                run.font.size = Pt(28)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 102, 204)
        
        elif line.startswith('## '):
            heading = doc.add_heading(line[3:], 1)
            for run in heading.runs:
                run.font.size = Pt(22)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 102, 204)
        
        elif line.startswith('### '):
            heading = doc.add_heading(line[4:], 2)
            for run in heading.runs:
                run.font.size = Pt(18)
                run.font.bold = True
        
        elif line.startswith('**') and line.endswith('**'):
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.add_run(line.replace('**', ''))
            run.bold = True
            run.font.size = Pt(14)
        
        elif line.startswith('[IMAGE:'):
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.add_run(line)
            run.italic = True
            run.font.color.rgb = RGBColor(128, 128, 128)
            run.font.size = Pt(11)
        
        elif line.startswith('|'):
            continue
        
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        
        elif line.strip() == '---':
            doc.add_paragraph()
        
        elif line.strip():
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.size = Pt(12)
        
        else:
            doc.add_paragraph()

    doc.save(docx_file)
    print(f"Successfully converted {md_file} to {docx_file}")

if __name__ == "__main__":
    if os.path.exists(input_file):
        convert_markdown_to_docx(input_file, output_file)
    else:
        print(f"Error: {input_file} not found!")